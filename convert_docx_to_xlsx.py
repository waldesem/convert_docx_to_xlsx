#pip install pandas==1.4.3 python-docx openpyxl
#Версия pandas выбрана из-за наличия в ней append
import sys
import os,os.path
from docx.api import Document
import pandas as pd
from pathlib import Path
#Входной каталог для перебора папок
input_dir = r'/home/MyProjects/Desktop3/'
#Перебираем рекурсивно каталоги; ищем файлы с таблицами Word
for subdir, dirs, files in os.walk(input_dir):
    for file in files:
        if file.endswith(".docx") and not file.lower().startswith("кронус"):
            in_file = os.path.join(subdir, file)
            print(in_file)
            output_file = file.split('.')[0]
            print(output_file)
            out_file = input_dir+output_file+'.xlsx'
            print(out_file)
            #out_file = r'C:\Users\ubuntu\Documents\Кандидаты\\'+output_file+'.xlsx'
            #Читаем документ, обрабатываем исключения файлов docx, которые созданы пустыми
            try:
                document = Document(in_file)
            except:
                document = Document()
            tables = document.tables
            #Проверяем, есть ли документы из которых удалено содержимое, если да - пропуускаем 
            all_paras = document.paragraphs
            if len(all_paras) < 1:
                break            
            df = pd.DataFrame()
            #Разбираем документ
            for table in document.tables:
                for row in table.rows:
                    text = [cell.text for cell in row.cells]
                    ### The frame.append method is deprecated and will be removed from pandas in a future version. Use pandas.concat instead.
                    df = df.append([text], ignore_index=True)
            #Количество Column зависит от количество колонок в файле заключения 2 или 3, скрываем заголовки и индексы
            if len(df.columns) == 4:
                df.columns = ["Column1", "Column2", "Column3", "Column4"]
                #Добавляем строку пути к папке проверки
                #df.loc[len(df)] = ["url", subdir]
                #Транспонируем таблицу
                #df=df.T
                df.to_excel(out_file, header = False, index = False)
            elif len(df.columns) == 3:
                df.columns = ["Column1", "Column2", "Column3"]
                #df.loc[len(df)] = ["url", "url", subdir]
                #df=df.T
                #df.drop(["Column2"]) Не удается выбросить вторую строку
                df.to_excel(out_file, header = False, index = False)
                #columns=["Column1", "Column3"] можно указать в скобках выше, но будут проблемы с транспонированием
            else:
                continue
        else:
            continue
        print(df)
